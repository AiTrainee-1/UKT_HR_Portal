"""
Resume Screening ML/scoring pipeline.
=====================================
Pure-function engine, independent of Django views — mirrors how
document_pdf.py is a standalone engine consumed by company_documents_views.py.

Lightweight-by-design stack (confirmed with the user): pdfplumber/python-docx
for text extraction, spaCy's small English model for name/city NER, rapidfuzz
for fuzzy skill matching, and scikit-learn TF-IDF + cosine similarity for a
lexical-overlap "semantic-ish" signal — deliberately NOT sentence-transformers/
PyTorch, which would add a ~300-700MB install and a runtime model-download
dependency unsuitable for this on-premise Windows deployment.

This is a hybrid rule-based + lexical-ML scorer, not a trained classifier —
there is no labeled hire/no-hire history to train one on, and every score
comes with a transparent, per-component breakdown (score_breakdown) rather
than being a black box, which matters for a hiring-adjacent decision.

screen_resume() is the single entry point used by BOTH the single-upload and
bulk-upload views, so there is exactly one scoring code path.
"""

import io
import re
from datetime import datetime
from decimal import Decimal

from rapidfuzz import fuzz


# ── spaCy singleton ─────────────────────────────────────────────────────────
# Loaded once at import time. If the model hasn't been downloaded
# (`python -m spacy download en_core_web_sm`), fail with one clear actionable
# message instead of a raw traceback burying the real cause.
_nlp = None
_nlp_error = None


def _get_nlp():
    global _nlp, _nlp_error
    if _nlp is not None:
        return _nlp
    if _nlp_error is not None:
        raise RuntimeError(_nlp_error)
    try:
        import spacy
        _nlp = spacy.load("en_core_web_sm")
        return _nlp
    except Exception as exc:  # OSError if the model isn't downloaded, ImportError if spacy isn't installed
        _nlp_error = (
            "Resume Screening needs the spaCy English model. Run this once inside "
            "the backend's .venv: python -m spacy download en_core_web_sm "
            f"(original error: {exc})"
        )
        raise RuntimeError(_nlp_error)


# ── Text extraction ─────────────────────────────────────────────────────────

def extract_text_from_pdf(file_obj) -> str:
    import pdfplumber
    text_parts = []
    with pdfplumber.open(file_obj) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n".join(text_parts)


def extract_text_from_docx(file_obj) -> str:
    import docx
    document = docx.Document(file_obj)
    parts = [p.text for p in document.paragraphs if p.text]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts)


def extract_text(file_obj, filename: str) -> str:
    ext = (filename.rsplit(".", 1)[-1] if "." in filename else "").lower()
    if ext == "pdf":
        return extract_text_from_pdf(file_obj)
    if ext in ("docx",):
        return extract_text_from_docx(file_obj)
    raise ValueError(
        f"Unsupported resume file type '.{ext}' — only .pdf and .docx are supported."
    )


# ── Identity extraction ─────────────────────────────────────────────────────

_EMAIL_RE = re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}")
# Indian mobile numbers: optional +91/91 prefix, 10 digits starting 6-9.
_PHONE_RE = re.compile(r"(?:\+?91[\s-]?)?\b[6-9]\d{9}\b")


def extract_email(text: str) -> str | None:
    m = _EMAIL_RE.search(text)
    return m.group(0) if m else None


def extract_phone(text: str) -> str | None:
    m = _PHONE_RE.search(text)
    if not m:
        return None
    return re.sub(r"[\s-]", "", m.group(0))


def extract_name(text: str, nlp=None) -> str | None:
    head = text[:600]
    try:
        nlp = nlp or _get_nlp()
        doc = nlp(head)
        for ent in doc.ents:
            if ent.label_ == "PERSON" and 2 <= len(ent.text.split()) <= 4:
                return ent.text.strip()
    except Exception:
        pass
    # Fallback: first non-empty line that isn't obviously a heading/contact line
    for line in head.splitlines():
        line = line.strip()
        if line and not _EMAIL_RE.search(line) and not _PHONE_RE.search(line) and len(line) <= 60:
            return line
    return None


def extract_city(text: str, known_cities: list[str], nlp=None) -> str | None:
    lower = text.lower()
    for city in known_cities:
        if city and city.lower() in lower:
            return city
    # Fallback: spaCy GPE entities, matched loosely against known_cities
    if known_cities:
        try:
            nlp = nlp or _get_nlp()
            doc = nlp(text[:1500])
            for ent in doc.ents:
                if ent.label_ == "GPE":
                    for city in known_cities:
                        if fuzz.ratio(ent.text.lower(), city.lower()) >= 85:
                            return city
        except Exception:
            pass
    return None


# ── Skills extraction ───────────────────────────────────────────────────────

def _build_vocabulary(rule_sets, field: str) -> list[str]:
    seen: dict[str, str] = {}
    for rs in rule_sets:
        for skill in (getattr(rs, field, None) or []):
            skill = str(skill).strip()
            if skill and skill.lower() not in seen:
                seen[skill.lower()] = skill
    return list(seen.values())


def build_skill_vocabulary(rule_sets) -> list[str]:
    """rule_sets: iterable of HiringRuleSet (or any object with .required_skills)."""
    return _build_vocabulary(rule_sets, "required_skills")


def build_soft_skill_vocabulary(rule_sets) -> list[str]:
    """rule_sets: iterable of HiringRuleSet (or any object with .soft_skills)."""
    return _build_vocabulary(rule_sets, "soft_skills")


def extract_skills(text: str, vocabulary: list[str], score_cutoff: int = 85) -> list[str]:
    lower = text.lower()
    matched = []
    for skill in vocabulary:
        skill_l = skill.lower()
        if len(skill_l) < 3:
            continue
        if skill_l in lower:
            matched.append(skill)
            continue
        if fuzz.partial_ratio(skill_l, lower) >= score_cutoff:
            matched.append(skill)
    return matched


# ── Experience extraction ───────────────────────────────────────────────────
#
# Evidence-only by design: a candidate's experience is reported ONLY when the
# resume gives an explicit, unambiguous statement of it — either a direct
# "X years of experience" line, or a sum of actual employment date ranges
# found within a detected Experience/Employment section. A bare "X years"
# anywhere in the document (a warranty period, a notice period, an unrelated
# certification blurb) is never trusted, and date ranges outside the
# experience section (e.g. an Education entry's "2015-2019") are never
# counted as job tenure. When neither form of evidence is present, the
# result is None (unknown) rather than a guess.

_DATE_RANGE_RE = re.compile(
    r"(19|20)\d{2}\s*(?:-|–|to)\s*((?:19|20)\d{2}|present|current|now)", re.IGNORECASE
)

# "Experience: 5 years" / "Total experience - 5+ years" / "Work experience: 5 yrs"
_EXPLICIT_EXPERIENCE_RE_A = re.compile(
    r"(?:total\s+)?(?:work\s+)?experience\s*[:\-]?\s*(\d+(?:\.\d+)?)\s*\+?\s*(?:years?|yrs?)\b",
    re.IGNORECASE,
)
# "5 years of experience" / "5+ years experience" / "3 years working experience"
_EXPLICIT_EXPERIENCE_RE_B = re.compile(
    r"(\d+(?:\.\d+)?)\s*\+?\s*(?:years?|yrs?)\s+(?:of\s+)?(?:total\s+|work(?:ing)?\s+)?experience\b",
    re.IGNORECASE,
)

_EXPERIENCE_SECTION_HEADER_RE = re.compile(
    r"^\s*(work\s+experience|professional\s+experience|employment\s+history|experience)\s*:?\s*$",
    re.IGNORECASE,
)
_SECTION_STOP_HEADER_RE = re.compile(
    r"^\s*(education|academic|qualification|skills|projects|certifications?|declaration|references)\b",
    re.IGNORECASE,
)


def _extract_experience_section(text: str) -> str | None:
    """Best-effort isolation of the resume's own Experience/Employment
    section, so education or certification date ranges elsewhere in the
    document are never mistaken for job tenure."""
    lines = text.splitlines()
    start = None
    for i, line in enumerate(lines):
        if _EXPERIENCE_SECTION_HEADER_RE.match(line.strip()):
            start = i + 1
            break
    if start is None:
        return None
    end = len(lines)
    for i in range(start, len(lines)):
        if _SECTION_STOP_HEADER_RE.match(lines[i].strip()):
            end = i
            break
    section = "\n".join(lines[start:end]).strip()
    return section or None


def extract_experience_years(text: str) -> float | None:
    # 1. An explicit, self-reported total — the highest-confidence evidence.
    explicit: list[float] = []
    for pattern in (_EXPLICIT_EXPERIENCE_RE_A, _EXPLICIT_EXPERIENCE_RE_B):
        for m in pattern.finditer(text):
            try:
                val = float(m.group(1))
                if 0 < val <= 45:
                    explicit.append(val)
            except ValueError:
                continue
    if explicit:
        return max(explicit)

    # 2. Sum of actual employment date ranges, but ONLY within a detected
    #    Experience/Employment section — never scanned across the whole resume.
    section = _extract_experience_section(text)
    if section:
        spans = []
        for m in _DATE_RANGE_RE.finditer(section):
            start_year = int(m.group(0)[:4])
            end_raw = m.group(2).lower()
            end_year = datetime.now().year if end_raw in ("present", "current", "now") else int(end_raw)
            if 1970 < start_year <= end_year <= datetime.now().year:
                spans.append(end_year - start_year)
        if spans:
            total = sum(spans)
            if 0 < total <= 45:
                return float(total)

    # No confident evidence either way — report unknown rather than guessing.
    return None


# ── Education extraction ────────────────────────────────────────────────────

# Ordered highest -> lowest; index is used as the ordinal rank.
EDUCATION_TIERS: list[tuple[str, list[str]]] = [
    ("phd", ["phd", "ph.d", "doctorate"]),
    ("postgraduate", ["m.tech", "mtech", "mba", "m.sc", "msc", "m.com", "postgraduate", "master of", "master's", "pg"]),
    ("graduate", ["b.tech", "btech", "b.e.", "b.sc", "bsc", "b.com", "bachelor", "graduate", "diploma in engineering", "ug"]),
    ("iti_diploma", ["iti", "diploma"]),
    ("12th", ["12th", "hsc", "higher secondary", "plus two", "+2"]),
    ("10th", ["10th", "ssc", "secondary school", "matriculation"]),
]
_TIER_ORDER = [key for key, _ in EDUCATION_TIERS]


def extract_education_level(text: str) -> str | None:
    lower = text.lower()
    for key, keywords in EDUCATION_TIERS:
        if any(kw in lower for kw in keywords):
            return key
    return None


_MORE_THAN_PREFIX_RE = re.compile(r"^\s*more\s+than\s+", re.IGNORECASE)


def education_meets_requirement(candidate_tier: str | None, required_text: str | None) -> bool:
    """
    required_text may be an exact tier ("B.Tech" -> candidate needs B.Tech or
    higher) or a strict "More than X" requirement ("More than B.Tech" ->
    candidate needs a tier strictly above B.Tech, e.g. a Master's).
    """
    if not required_text or not required_text.strip():
        return True  # no requirement configured -> nothing to fail

    strict = bool(_MORE_THAN_PREFIX_RE.match(required_text))
    base_text = _MORE_THAN_PREFIX_RE.sub("", required_text) if strict else required_text

    required_tier = extract_education_level(base_text)
    if required_tier is None:
        return True  # requirement text didn't resolve to a recognizable tier -> don't penalize
    if candidate_tier is None:
        return False

    candidate_rank = _TIER_ORDER.index(candidate_tier)
    required_rank = _TIER_ORDER.index(required_tier)
    return candidate_rank < required_rank if strict else candidate_rank <= required_rank


# ── TF-IDF similarity ───────────────────────────────────────────────────────

def compute_tfidf_similarity(resume_text: str, rule_set_text: str) -> float:
    if not resume_text.strip() or not rule_set_text.strip():
        return 0.0
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.metrics.pairwise import cosine_similarity
    try:
        vectorizer = TfidfVectorizer(stop_words="english")
        matrix = vectorizer.fit_transform([resume_text, rule_set_text])
        if matrix.shape[1] == 0:
            return 0.0
        sim = cosine_similarity(matrix[0:1], matrix[1:2])[0][0]
        return max(0.0, min(1.0, float(sim)))
    except ValueError:
        return 0.0


# ── Final weighted scoring ──────────────────────────────────────────────────

WEIGHTS = {"skills": 35, "softSkills": 5, "similarity": 25, "experience": 20, "education": 10, "location": 5}


def _rule_set_text(rule_set) -> str:
    parts = [
        " ".join(rule_set.required_skills or []),
        " ".join(rule_set.soft_skills or []),
        rule_set.education_qualification or "",
        rule_set.other_requirements or "",
    ]
    return " ".join(p for p in parts if p)


def score_candidate(resume_text: str, rule_set, extracted: dict) -> dict:
    required_skills = rule_set.required_skills or []
    matched_skills = extracted.get("extracted_skills") or []
    skills_score = (
        (len(matched_skills) / len(required_skills)) * WEIGHTS["skills"]
        if required_skills else WEIGHTS["skills"]
    )
    skills_score = min(skills_score, WEIGHTS["skills"])

    soft_skills = rule_set.soft_skills or []
    matched_soft_skills = extracted.get("extracted_soft_skills") or []
    soft_skills_score = (
        (len(matched_soft_skills) / len(soft_skills)) * WEIGHTS["softSkills"]
        if soft_skills else WEIGHTS["softSkills"]
    )
    soft_skills_score = min(soft_skills_score, WEIGHTS["softSkills"])

    similarity_raw = compute_tfidf_similarity(resume_text, _rule_set_text(rule_set))
    similarity_score = similarity_raw * WEIGHTS["similarity"]

    required_years = float(rule_set.min_experience_years or 0)
    candidate_years = extracted.get("extracted_experience_years")
    if required_years <= 0:
        experience_score = WEIGHTS["experience"]
    elif candidate_years is None:
        experience_score = 0.0
    elif candidate_years >= required_years:
        experience_score = WEIGHTS["experience"]
    else:
        experience_score = (candidate_years / required_years) * WEIGHTS["experience"]

    candidate_tier = extracted.get("extracted_education")
    education_meets = education_meets_requirement(candidate_tier, rule_set.education_qualification)
    education_score = WEIGHTS["education"] if education_meets else 0.0

    preferred_city = (rule_set.preferred_city or "").strip()
    candidate_city = extracted.get("city")
    location_meets = (
        not preferred_city
        or (candidate_city is not None and fuzz.ratio(candidate_city.lower(), preferred_city.lower()) >= 80)
    )
    location_score = WEIGHTS["location"] if location_meets else 0.0

    total = round(
        skills_score + soft_skills_score + similarity_score + experience_score + education_score + location_score, 2
    )

    return {
        "total": total,
        "components": {
            "skills": {
                "score": round(skills_score, 2), "weight": WEIGHTS["skills"],
                "matched": matched_skills, "missing": [s for s in required_skills if s not in matched_skills],
            },
            "softSkills": {
                "score": round(soft_skills_score, 2), "weight": WEIGHTS["softSkills"],
                "matched": matched_soft_skills, "missing": [s for s in soft_skills if s not in matched_soft_skills],
            },
            "similarity": {
                "score": round(similarity_score, 2), "weight": WEIGHTS["similarity"],
                "raw_cosine": round(similarity_raw, 4),
            },
            "experience": {
                "score": round(experience_score, 2), "weight": WEIGHTS["experience"],
                "required": required_years, "extracted": candidate_years,
            },
            "education": {
                "score": round(education_score, 2), "weight": WEIGHTS["education"],
                "required": rule_set.education_qualification, "extracted": candidate_tier, "meets": education_meets,
            },
            "location": {
                "score": round(location_score, 2), "weight": WEIGHTS["location"],
                "preferred": preferred_city or None, "extracted": candidate_city, "meets": location_meets,
            },
        },
    }


# ── Orchestration ────────────────────────────────────────────────────────────

def screen_resume(
    file_bytes: bytes, filename: str, rule_set,
    vocabulary: list[str], soft_skill_vocabulary: list[str], known_cities: list[str],
) -> dict:
    """
    Runs the full pipeline for one resume against one rule set.
    Returns a flat dict of fields ready to populate a ScreeningCandidate.
    Raises ValueError for unsupported file types or unreadable files —
    callers should catch this per-file so one bad resume doesn't abort a batch.
    """
    text = extract_text(io.BytesIO(file_bytes), filename)
    if not text.strip():
        raise ValueError("Could not extract any text from this file (it may be a scanned image with no text layer).")

    nlp = None
    try:
        nlp = _get_nlp()
    except RuntimeError:
        pass  # extract_name/extract_city fall back to regex/substring matching without it

    extracted_skills = extract_skills(text, vocabulary)
    extracted_soft_skills = extract_skills(text, soft_skill_vocabulary)
    extracted_experience_years = extract_experience_years(text)
    extracted_education = extract_education_level(text)
    candidate_name = extract_name(text, nlp)
    city = extract_city(text, known_cities, nlp)
    email = extract_email(text)
    phone = extract_phone(text)

    extracted = {
        "extracted_skills": extracted_skills,
        "extracted_soft_skills": extracted_soft_skills,
        "extracted_experience_years": extracted_experience_years,
        "extracted_education": extracted_education,
        "city": city,
    }
    breakdown = score_candidate(text, rule_set, extracted)

    return {
        "candidate_name": candidate_name,
        "email": email,
        "phone": phone,
        "city": city,
        "extracted_skills": extracted_skills,
        "extracted_soft_skills": extracted_soft_skills,
        "extracted_experience_years": (
            Decimal(str(extracted_experience_years)) if extracted_experience_years is not None else None
        ),
        "extracted_education": extracted_education,
        "raw_text_excerpt": text[:2000],
        "match_score": Decimal(str(breakdown["total"])),
        "score_breakdown": breakdown,
    }
